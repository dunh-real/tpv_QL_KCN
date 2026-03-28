import io
import base64
import logging
from pathlib import Path
from typing import List, Optional

from docx import Document as DocxDocument
from pdf2image import convert_from_path
from langchain_core.documents import Document
from langchain_ollama import ChatOllama
from langchain_core.messages import HumanMessage

from src.core.config import settings
from src.core.logger import get_logger

logger = get_logger(__name__)


class StrictDocumentExtractor:
    """
    Trích xuất tài liệu chuẩn RAG với quy tắc định tuyến cứng:
    - Mặc định dùng GLM-OCR (Vision LLM) cho mọi file PDF để lấy Markdown.
    - Dùng thư viện native cho DOCX và TXT để tối ưu tốc độ.
    """
    def __init__(self, model_name: str = settings.OCR_MODEL, temperature: float = 0.0):
        self.model_name = model_name
        try:

            self.llm = ChatOllama(model=self.model_name, temperature=temperature, base_url=settings.OLLAMA_BASE_URL)
            logger.info(f"Đã khởi tạo Extractor với model: {self.model_name}")
        except Exception as e:
            logger.error(f"Lỗi khởi tạo LLM: {e}")
            raise

    def _get_metadata(self, file_path: Path) -> dict:
        """Tạo metadata"""
        return {
            "source": str(file_path.name),
            "extension": file_path.suffix.lower(),
        }

    def _process_txt(self, file_path: Path, metadata: dict) -> List[Document]:
        """Đọc trực tiếp file .txt"""
        logger.info(f"Đang xử lý file TXT: {file_path.name}")
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read().strip()
            
            if not content:
                return []
            
            page_meta = metadata.copy()
            page_meta["page"] = 1
            return [Document(page_content=content, metadata=page_meta)]
        except Exception as e:
            logger.error(f"Lỗi đọc TXT {file_path.name}: {e}")
            raise

    def _process_docx(self, file_path: Path, metadata: dict) -> List[Document]:
        """Đọc trực tiếp file .docx"""
        logger.info(f"Đang xử lý file DOCX: {file_path.name}")
        try:
            doc = DocxDocument(str(file_path))
            content = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            
            if not content:
                return []
                
            page_meta = metadata.copy()
            page_meta["page"] = 1
            return [Document(page_content=content, metadata=page_meta)]
        except Exception as e:
            logger.error(f"Lỗi đọc DOCX {file_path.name}: {e}")
            raise

    def _image_to_base64(self, pil_image) -> str:
        """Chuẩn hóa ảnh và chuyển Base64 ."""
        if pil_image.mode != "RGB":
            pil_image = pil_image.convert("RGB")
            
        max_size = 1024
        if max(pil_image.size) > max_size:
            pil_image.thumbnail((max_size, max_size))
            logger.info(f"Đã hạ độ phân giải ảnh xuống: {pil_image.size}")

        buffered = io.BytesIO()
        pil_image.save(buffered, format="PNG")
        return base64.b64encode(buffered.getvalue()).decode("utf-8")

    def _process_pdf_with_vision(self, file_path: Path, metadata: dict) -> List[Document]:
        """
        Xử lý PDF: Cắt thành ảnh và ép GLM-OCR đọc ra Markdown.
        Lưu ý trên macOS: Yêu cầu đã cài poppler (brew install poppler).
        """
        logger.info(f"Đang kích hoạt GLM-OCR cho PDF: {file_path.name}")
        documents = []
        try:
            pages = convert_from_path(str(file_path))
            total_pages = len(pages)
            
            system_prompt = (
                "Bạn là hệ thống OCR chuyên dụng cho RAG. Hãy chuyển đổi hình ảnh này thành Markdown. "
                "Bảo toàn các thẻ tiêu đề (#), danh sách (*), và bóc tách bảng biểu thành format bảng Markdown. "
                "CHỈ trả về nội dung Markdown, KHÔNG giải thích gì thêm."
            )

            for i, page_image in enumerate(pages):
                page_num = i + 1
                logger.info(f"Đang OCR bằng AI trang {page_num}/{total_pages}...")
                
                base64_image = self._image_to_base64(page_image)
                
                messages = [
                    HumanMessage(
                        content=[
                            {"type": "text", "text": system_prompt},
                            {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{base64_image}"}},
                        ]
                    )
                ]

                response = self.llm.invoke(messages)
                content = response.content.strip()

                if content:
                    page_meta = metadata.copy()
                    page_meta["page"] = page_num
                    documents.append(Document(page_content=content, metadata=page_meta))

                del page_image
                del base64_image

            return documents
        except Exception as e:
            logger.error(f"Lỗi OCR PDF {file_path.name}: {e}")
            raise

    def extract(self, file_path: str | Path) -> Optional[List[Document]]:
        """
        API chính của class. Tự động xem đuôi file để quyết định gọi hàm nào.
        """
        path = Path(file_path)
        if not path.exists() or not path.is_file():
            logger.error(f"File không tồn tại: {path}")
            raise FileNotFoundError(f"Không tìm thấy file: {path}")

        metadata = self._get_metadata(path)
        ext = metadata["extension"]

        if ext == '.pdf':
            return self._process_pdf_with_vision(path, metadata)
        elif ext == '.docx':
            return self._process_docx(path, metadata)
        elif ext == '.txt':
            return self._process_txt(path, metadata)
        else:
            raise ValueError(f"Định dạng {ext} không được hỗ trợ. Chỉ nhận .pdf, .docx, .txt")